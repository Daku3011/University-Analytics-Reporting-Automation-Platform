"""Word (DOCX) export of the Annual Portfolio Report (#5)."""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

BRAND = RGBColor(0x1F, 0x3A, 0x5F)
MUTED = RGBColor(0x66, 0x66, 0x66)


def _fmt(value):
    return f"{value or 0:,}"


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND
    return h


def _table(doc, headers, rows, number_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            text = _fmt(value) if i in number_cols else str(value)
            cells[i].text = text
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
                    if i in number_cols:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return table


def build_portfolio_docx(context):
    """Render the portfolio context into a .docx document, returned as bytes."""
    chapters = context['chapters']
    ex = chapters['executive_summary']
    doc = Document()

    # ── Title block ─────────────────────────────────────────────────
    title = doc.add_heading(f"{context['college_name']}", level=0)
    for run in title.runs:
        run.font.color.rgb = BRAND
    subtitle = doc.add_paragraph(
        f"Annual Portfolio Report — {ex['year']} · Sarvajanik University")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.color.rgb = MUTED
        run.font.size = Pt(11)

    # ── Chapter 1: Executive summary ────────────────────────────────
    _heading(doc, '1. Executive Summary')
    doc.add_paragraph(
        f"This report consolidates {ex['year']} communications performance for "
        f"{context['college_name']}: social media output and reach, events, media "
        f"coverage, press releases and progress against KPI targets.")
    _table(
        doc,
        ['Metric', str(ex['prev_year']), str(ex['year']), 'Change'],
        [[m['label'], m['previous'], m['current'],
          '—' if m['change'] is None else f"{m['change']:+.1f}%"] for m in ex['yoy']],
        number_cols=(1, 2),
    )
    facts = [
        ('Events held', str(ex['event_count'])),
        ('Media coverage items', str(ex['media_count'])),
        ('Press releases', str(ex['press_release_count'])),
        ('Top posts', str(ex['top_post_count'])),
        ('Months reported', str(ex['months_reported'])),
        ('KPI targets met', f"{ex['kpi_targets_met']} of {ex['kpi_target_count']}"),
        ('Average KPI attainment',
         '—' if ex['kpi_avg_attainment'] is None else f"{ex['kpi_avg_attainment']}%"),
    ]
    _table(doc, ['Highlights', 'Value'], facts)

    # ── Chapter 2: Social media ────────────────────────────────────
    sm = chapters['social_media']
    _heading(doc, '2. Social Media Performance')
    if sm['monthly']:
        _table(
            doc,
            ['Month', 'Total Views', 'Total Reach', 'Followers Gained', 'Reels', 'Graphics'],
            [[row['record'].get_month_display(), row['record'].total_views,
              row['record'].total_reach, row['record'].followers_gained,
              row['record'].reels_count, row['record'].graphics_count]
             for row in sm['monthly']],
            number_cols=(1, 2, 3, 4, 5),
        )
    else:
        doc.add_paragraph('No monthly analytics were recorded for this year.')
    if sm['top_posts']:
        _heading(doc, 'Top Posts', level=2)
        _table(
            doc,
            ['Month', 'Platform', 'Caption', 'Views', 'Likes'],
            [[p.get_month_display(), p.platform.title(), (p.caption or '—')[:80],
              p.views, p.likes] for p in sm['top_posts']],
            number_cols=(3, 4),
        )

    # ── Chapter 3: Events ──────────────────────────────────────────
    ev = chapters['events']
    _heading(doc, '3. Events')
    if ev['events']:
        category_line = ', '.join(
            f"{name.replace('_', ' ').title()} ×{count}" for name, count in ev['categories'])
        doc.add_paragraph(f"{ev['total']} events recorded — by category: {category_line}.")
        _table(
            doc,
            ['Date', 'Title', 'Category'],
            [[e.date.strftime('%d %b %Y') if e.date else '—', e.title,
              (e.category or 'other').replace('_', ' ').title()] for e in ev['events']],
        )
    else:
        doc.add_paragraph('No events were recorded for this year.')

    # ── Chapter 4: Media coverage ──────────────────────────────────
    mc = chapters['media_coverage']
    _heading(doc, '4. Media Coverage')
    if mc['newspapers']:
        _heading(doc, 'Newspapers', level=2)
        _table(
            doc,
            ['Publication', 'Date', 'Headline'],
            [[n.publication, n.date.strftime('%d %b %Y') if n.date else '—',
              n.headline or '—'] for n in mc['newspapers']],
        )
    if mc['channels']:
        _heading(doc, 'TV / Channels', level=2)
        _table(
            doc,
            ['Channel', 'Month', 'Programme / Platform'],
            [[c_.channel_name, f"{c_.month}/{c_.year}", c_.platform or '—']
             for c_ in mc['channels']],
        )
    if not mc['newspapers'] and not mc['channels']:
        doc.add_paragraph('No media coverage was recorded for this year.')

    # ── Chapter 5: Press releases ──────────────────────────────────
    pr = chapters['press_releases']
    _heading(doc, '5. Press Releases')
    if pr['releases']:
        doc.add_paragraph(
            f"{pr['total']} releases issued with {pr['total_placements']} total placements.")
        _table(
            doc,
            ['Date', 'Title', 'Placements', 'Reach'],
            [[p.date_submitted.strftime('%d %b %Y') if p.date_submitted else '—',
              p.title, p.placements, str(p.potential_reach or '—')]
             for p in pr['releases']],
            number_cols=(2,),
        )
    else:
        doc.add_paragraph('No press releases were issued this year.')

    # ── Chapter 6: KPI performance ─────────────────────────────────
    kpi = chapters['kpi_performance']['rows']
    _heading(doc, '6. KPI Performance')
    if kpi:
        _table(
            doc,
            ['Scope', 'Metric', 'Target', 'Actual', 'Gap', 'Achievement'],
            [[str(r_['scope']), r_['metric'], r_['target'], r_['actual'], r_['gap'],
              f"{r_['achievement']}% {'✔' if r_['on_track'] else '✖'}"]
             for r_ in kpi],
            number_cols=(2, 3, 4),
        )
    else:
        doc.add_paragraph('No KPI targets were defined for this year.')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
